import os
from docx import Document
import pandas as pd

# ler a tabela do Excel
tabela_excel = pd.read_excel('seu_caminho')

def preencher_contrato(template_path, dados_cliente, nome_arquivo):

    document = Document(template_path)

    for paragraph in document.paragraphs:
        for chave, valor in dados_cliente.items():
            if chave in paragraph.text:
                paragraph.text = paragraph.text.replace(chave, valor)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for chave, valor in dados_cliente.items():
                        if chave in paragraph.text:
                            paragraph.text = paragraph.text.replace(chave, valor)

    # pasta onde os contratos serão salvos
    pasta_contratos = 'contratos_prontos'

    # verifica se a pasta existe, se não, cria a pasta
    if not os.path.exists(pasta_contratos):
        os.makedirs(pasta_contratos)

    # salvar o contrato preenchido na pasta contratos_prontos
    document.save(os.path.join(pasta_contratos, f'{nome_arquivo}.docx'))

# iterar sobre cada linha da tabela do Excel
for index, linha in tabela_excel.iterrows():
    # Criar um dicionário com os dados do cliente
    dados_cliente = {
        "[razao_social]": str(linha['razao_social']) if not pd.isnull(linha['razao_social']) else "",
        "[cnpj]": str(linha['cnpj']) if not pd.isnull(linha['cnpj']) else "",
        "[email]": str(linha['email']) if not pd.isnull(linha['email']) else "",
        "[phone1]": str(linha['phone1']) if not pd.isnull(linha['phone1']) else "",
        "[phone2]": str(linha['phone2']) if not pd.isnull(linha['phone2']) else "",
        "[adress]": str(linha['adress']) if not pd.isnull(linha['adress']) else "",
        "[bairro]": str(linha['bairro']) if not pd.isnull(linha['bairro']) else "",
        "[complement]": str(linha['complement']) if not pd.isnull(linha['complement']) else "",
        "[cidade]": str(linha['cidade']) if not pd.isnull(linha['cidade']) else "",
        "[uf]": "GO",
        "[cep]": str(linha['cep']) if not pd.isnull(linha['cep']) else "",
        "[data]": "ADICIONAR DATA"
    }
    
    # caminho para o template do contrato
    template_path = "seu_caminho"

    # nome do arquivo com base no código do cliente
    nome_arquivo = f"contrato_{linha['razao_social']}"
    
    # preencher e salvar o contrato
    preencher_contrato(template_path, dados_cliente, nome_arquivo)
